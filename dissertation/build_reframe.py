# Reframe the dissertation around RQ-D / RQ-A / RQ-O, editing a COPY of the
# WITH_FIGURES docx (figures preserved). python-docx (no docxjs skill exists).
import shutil
import docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r'Does_Observability_Matter_Dissertation.WITH_FIGURES.docx'
DST = r'Does_Observability_Matter_Dissertation.REFRAMED.docx'
shutil.copyfile(SRC, DST)
doc = docx.Document(DST)
ps = doc.paragraphs


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text)


def sname(p):
    try:
        return p.style.name or ''
    except Exception:
        return ''


# -------------------- new front-matter text --------------------
ABSTRACT = (
 "This dissertation asks whether observability matters for real-time anomaly detection in "
 "cloud-native systems, and answers the question along three operational dimensions that a "
 "foundational study of observability completeness opens up. That foundational study establishes "
 "that detection quality rises monotonically with telemetry richness and that distributed traces "
 "are the single most valuable signal, lifting root-cause localisation to perfection. Building on "
 "this, the dissertation poses three primary research questions about the operational value of that "
 "observability. First, on timeliness: with a gradual fault onset and a defined service-level "
 "objective, a proactive forecasting detector warns earlier than a reactive classifier—a "
 "two-window versus one-window median lead, and an eighty-nine versus sixty-six per cent "
 "early-warning rate—at the cost of a modest rise in false alarms. Second, on robustness: "
 "detection degrades gracefully under reduced trace sampling and noise, but a modality-fragility "
 "analysis shows that traces, the most valuable signal, are also the most fragile, their loss "
 "costing far more detection accuracy than the loss of logs or events. Third, on cost-efficiency: "
 "framing the choice of observability configuration as a multi-objective optimisation reveals that "
 "full Metrics-Events-Logs-Traces telemetry is cost-optimal only at the extreme; for almost any "
 "budget a traces-centric, sampled configuration dominates, and a metrics-plus-traces configuration "
 "attains within a fraction of a per cent of the best detection at roughly a quarter less cost. "
 "Taken together, the three studies reframe the guiding question: observability matters not merely "
 "for whether failures are detected, but for how early, how robustly, and how economically they are "
 "detected. All experiments are implemented on a controlled, fault-injected testbed and are fully "
 "reproducible.")

AIM = (
 "The aim of this dissertation is to determine, through controlled experiment, the operational value "
 "of observability for real-time anomaly detection in cloud-native systems—not only whether "
 "richer telemetry improves detection, but how that detection performs along the three axes that "
 "matter in production: its timeliness, its robustness to imperfect telemetry, and its "
 "cost-efficiency.")

LEAD = (
 "A foundational study first establishes the value of observability completeness (Chapters 4–5), "
 "finding that detection quality rises monotonically with telemetry richness and that distributed "
 "traces are the decisive signal. Building on that finding, the dissertation pursues three primary "
 "research questions, each operationalised as a concrete experiment in the extended empirical study "
 "(Chapter 8):")

RQD = (
 "RQ-D (Timeliness): How early, before a service-level objective is breached, can observability-driven "
 "detection raise an alarm, and how does a proactive forecasting approach compare with a reactive "
 "classifier on the trade-off between detection accuracy and lead time?")
RQA = (
 "RQ-A (Robustness): How robust is observability-driven detection to realistic telemetry "
 "degradation—reduced trace sampling, missing modalities, and noisy signals—and which "
 "telemetry signals are the most fragile?")
RQO = (
 "RQ-O (Cost-efficiency): Given a telemetry budget, which observability configuration—modality "
 "mix and trace-sampling rate—maximises detection quality per unit cost, and is full MELT "
 "telemetry actually cost-optimal?")

# -------------------- apply front-matter edits --------------------
# Abstract: replace first body paragraph after the heading, clear the rest until next H1.
ab = next(i for i, p in enumerate(ps) if sname(p).startswith('Heading 1')
          and p.text.strip().lower() == 'abstract')
end = next((j for j in range(ab + 1, len(ps)) if sname(ps[j]).startswith('Heading 1')), len(ps))
body = [j for j in range(ab + 1, end) if ps[j].text.strip()]
set_text(ps[body[0]], ABSTRACT)
for j in body[1:]:
    set_text(ps[j], '')

set_text(ps[34], AIM)
set_text(ps[35], LEAD)
set_text(ps[36], RQD)
set_text(ps[37], RQA)
set_text(ps[38], RQO)

# -------------------- new Chapter 8 (inserted before References) --------------------
CH = [
 ("Heading 1", "Chapter 8: Extended Empirical Study — Early Detection, Robustness and Cost-Efficiency"),
 ("Body",
  "The foundational study (Chapters 4–5) established that observability completeness improves "
  "detection and that traces are the decisive signal. That study, however, measured only detection "
  "accuracy on a clean, stationary, fully-instrumented stream. This chapter extends the investigation "
  "to the three operational dimensions that determine whether a detector is useful in "
  "production—timeliness, robustness and cost-efficiency—each through a controlled experiment "
  "implemented on the same testbed. The experiments are deliberately built on the foundational "
  "finding: because traces are the most valuable signal, the questions of how early they warn, how "
  "fragile they are, and how cheaply they can be collected are precisely the questions that matter."),

 ("Heading 2", "8.1 RQ-D: How Early Can Detection Warn Before an SLO Breach?"),
 ("Body",
  "To make ‘early’ measurable, faults are given a gradual onset: within a fault episode the "
  "fault intensity ramps from zero to full, so the tail latency climbs and eventually crosses a "
  "defined service-level objective at a well-identified breach window. Two detectors are compared by "
  "their lead time, the number of windows between the alarm and the breach. The first is the reactive "
  "classifier used throughout the foundational study, which flags the current window once it appears "
  "anomalous. The second is a proactive forecaster that fits the recent latency trend and raises an "
  "alarm when its forecast, several windows ahead, is projected to cross the objective. Normal "
  "episodes, in which no breach occurs, measure each approach’s false-alarm rate."),
 ("Body",
  "Across three hundred episodes, with the median breach at the eleventh window, the reactive detector "
  "achieved a median lead of one window, warning before the breach in sixty-six per cent of episodes "
  "and never raising a false alarm. The proactive forecaster achieved a median lead of two windows and "
  "an early-warning rate of eighty-nine per cent, at the cost of a nine per cent false-alarm rate on "
  "normal episodes."),
 ("Body",
  "The result establishes a clear and tunable trade-off: forecasting buys earlier warning, but "
  "earliness is paid for in false alarms. For a service whose error budget tolerates occasional "
  "spurious alerts, the additional lead time is valuable, because it converts detection from a record "
  "of failure into a chance to prevent it; for a latency-critical service intolerant of alert noise, "
  "the reactive approach’s perfect precision may be preferred. Timeliness is therefore a design "
  "axis to be chosen deliberately, not a fixed property of the detector."),

 ("Heading 2", "8.2 RQ-A: How Robust Is Detection to Telemetry Degradation?"),
 ("Body",
  "Robustness is examined by training a detector on clean, full-telemetry data and then evaluating it "
  "on telemetry degraded in three realistic ways, each swept across severity levels. Trace sampling "
  "scales the trace-derived signals by the sampling rate and drops the rare originating error-spans "
  "when they are not sampled; multiplicative noise perturbs every telemetry field; and modality "
  "dropout removes an entire pillar. The drop in detection accuracy from the clean baseline measures "
  "the fragility of detection to each form of degradation, and the per-pillar dropout sweep ranks "
  "which signal the detector most depends upon."),
 ("Body",
  "On a clean full-telemetry baseline the detector achieved an F1 of 0.99. Detection degraded "
  "gracefully under trace sampling, falling from 0.99 at full sampling to 0.81 at a twentieth of the "
  "traces, and was robust to noise up to a moderate level before declining. The modality-fragility "
  "ranking was decisive: removing traces cost 0.18 in F1, whereas removing logs or events cost "
  "essentially nothing."),
 ("Body",
  "The finding sharpens the foundational result in an operationally important way. Traces are "
  "simultaneously the most valuable signal and the most fragile: the very pillar whose addition most "
  "improved detection is the one whose degradation most harms it. A detection pipeline that depends on "
  "traces therefore inherits a dependency on the reliability and sampling fidelity of the tracing "
  "backend, which is the most expensive and most aggressively sampled pillar in practice. Robustness "
  "and value are, for traces, two sides of the same coin."),

 ("Heading 2", "8.3 RQ-O: What Is the Cost-Optimal Observability Configuration?"),
 ("Body",
  "The cost-efficiency question is posed as a multi-objective optimisation over twenty observability "
  "configurations, each a combination of a modality subset—metrics are always collected, while "
  "logs, traces and events are each optional—and, where traces are collected, a sampling rate. "
  "Each configuration is assigned a telemetry cost, with traces weighted as the most expensive pillar "
  "and that cost scaled by the sampling rate, and a detection F1 obtained by training the detector on "
  "exactly that configuration. The F1-versus-cost Pareto front identifies which configurations are not "
  "dominated by a cheaper one of equal or better accuracy."),
 ("Body",
  "Full Metrics-Events-Logs-Traces telemetry achieved the highest detection F1 of 0.99 and lay on the "
  "Pareto front, but only at the front’s extreme, at the highest cost. The knee of the front was a "
  "metrics-plus-traces configuration, which attained an F1 of 0.99—within three thousandths of the "
  "best—at roughly a quarter less cost. Logs never reached the efficient frontier: every "
  "configuration that added logs was dominated. Trace sampling proved an almost free economy: a "
  "metrics-traces-events configuration at a tenth of the trace volume reached an F1 of 0.96 at under a "
  "quarter of the cost of full telemetry."),
 ("Body",
  "The result qualifies the foundational ‘more is better’ conclusion precisely. More "
  "observability is better only at the extreme high-accuracy end; for almost any realistic budget a "
  "traces-centric, sampled configuration dominates full telemetry, and logs—despite their "
  "cost—add little once traces are present. The minimum-viable observability for detection is "
  "metrics with sampled traces, and the substantial expense of full-fidelity, all-pillar telemetry is "
  "justified only when the last fraction of a per cent of detection accuracy is worth several times the "
  "cost."),

 ("Heading 2", "8.4 Synthesis"),
 ("Body",
  "The three studies converge on a single, sharper answer to the dissertation’s guiding question. "
  "Observability matters for anomaly detection not as a binary—detected or not—but along three "
  "operational dimensions, and on each the decisive role of traces, established in the foundational "
  "study, reappears. Traces give the earliest warning and the most accurate localisation; they are the "
  "most fragile signal under degradation; and they are the pillar around which the cost-optimal "
  "configuration is built. The unifying recommendation is to instrument richly enough to capture "
  "traces, to sample them intelligently rather than collect them in full, to monitor the health of the "
  "tracing pipeline as a first-class concern, and to add a forecasting layer where lead time is worth a "
  "modest rise in false alarms. Observability matters; what this dissertation adds is a measured "
  "account of how early, how robustly, and how economically it can be made to matter."),
]

# anchor: the References heading (insert chapter before it); fallback = end of doc
anchor = next((p for p in doc.paragraphs if sname(p).startswith('Heading 1')
               and p.text.strip().lower() in ('references', 'bibliography')), None)


def ins_before(anchor, style, text):
    el = OxmlElement('w:p')
    anchor._p.addprevious(el)
    para = Paragraph(el, anchor._parent)
    try:
        para.style = doc.styles['Heading 1' if style == 'Heading 1'
                                else 'Heading 2' if style == 'Heading 2' else 'Normal']
    except Exception:
        pass
    if text:
        para.add_run(text)
    return para


if anchor is not None:
    for style, text in CH:
        ins_before(anchor, style, text)
else:
    for style, text in CH:
        if style == 'Heading 1':
            doc.add_heading(text, level=1)
        elif style == 'Heading 2':
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

doc.save(DST)
print("saved", DST)
print("abstract para set; RQ-D/A/O set at 36-38; chapter 8 inserted before",
      "References" if anchor is not None else "(appended at end)")
PY = None
