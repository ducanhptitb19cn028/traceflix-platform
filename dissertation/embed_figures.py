# Embed the three AIOps-generated figures into Chapter 8 of the reframed docx,
# each placed after its results paragraph, with a caption.
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

F = "Does_Observability_Matter_Dissertation.REFRAMED.docx"
OUT = "Does_Observability_Matter_Dissertation.FINAL.docx"
IMG = "D:/ResearchWithDrSatish/traceflix-platform/aiops/data/results/figures/"
d = docx.Document(F)


def insert_after(anchor, img, width, caption):
    # image paragraph, centred, right after the anchor
    ip = OxmlElement("w:p"); anchor._p.addnext(ip)
    img_para = Paragraph(ip, anchor._parent)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(img, width=Inches(width))
    # caption paragraph, right after the image
    cp = OxmlElement("w:p"); ip.addnext(cp)
    cap = Paragraph(cp, anchor._parent)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
    return img_para


SPEC = [
    ("Across three hundred episodes", IMG + "fig_rqD_leadtime.png", 4.4,
     "Figure 8.1: Lead-time distributions for the reactive and proactive detectors over fault "
     "episodes (positive = warning before the SLO breach). The forecaster's distribution is "
     "shifted earlier, giving a longer median lead."),
    ("On a clean full-telemetry baseline", IMG + "fig_rqA_robustness.png", 6.0,
     "Figure 8.2: Detection F1 under trace sampling (left) and the fragility of detection to "
     "dropping each telemetry pillar (right). Detection degrades gracefully with sampling, but "
     "traces are by far the most fragile signal."),
    ("Full Metrics-Events-Logs-Traces telemetry achieved the highest", IMG + "fig_rqO_pareto.png", 5.2,
     "Figure 8.3: Detection F1 versus telemetry cost across observability configurations; the "
     "Pareto front is highlighted. Full MELT is cost-optimal only at the extreme, while the knee "
     "of the front is a metrics-plus-traces configuration."),
]

n = 0
for needle, img, w, cap in SPEC:
    anchor = next((p for p in d.paragraphs if needle in p.text), None)
    if anchor is None:
        print("ANCHOR NOT FOUND:", needle[:40]); continue
    insert_after(anchor, img, w, cap)
    n += 1
    print("embedded after:", needle[:40])

d.save(OUT)
d2 = docx.Document(OUT)
print(f"\nembedded {n}/3 figures -> {OUT}; inline images now: {len(d2.inline_shapes)}")
