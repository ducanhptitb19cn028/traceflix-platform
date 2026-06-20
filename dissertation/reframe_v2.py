#!/usr/bin/env python3
"""Reframe the dissertation around the three PRIMARY RQs (RQ-D/A/O).

- keep old 5.2-5.4 as relabelled FOUNDATIONAL findings (no longer 'RQ1/2/3')
- promote Chapter 8's RQ-D/A/O sections (with figures) into Chapter 5 as 5.5-5.8
- delete the Chapter 8 shell; renumber the old Discussion to 5.9
- reword every residual 'first/second/third research question' to foundational
- embed the six draw.io diagrams (Fig 3.2, 4.1-4.5)
- weave LBU-Harvard in-text citations; merge all refs into one alphabetical list
Reads FINAL.docx, writes RQREFRAMED.docx (new file -> no Word lock).
"""
import json
from pathlib import Path

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

HERE = Path("D:/ResearchWithDrSatish/traceflix-platform/dissertation")
DIAG = HERE / "diagrams"
SRC = HERE / "Does_Observability_Matter_Dissertation.FINAL.docx"
OUT = HERE / "Does_Observability_Matter_Dissertation.RQREFRAMED.docx"
REFS = json.loads((HERE / "lbu_refs.json").read_text(encoding="utf-8"))

d = docx.Document(SRC)
P = d.paragraphs
BODY = P[0]._parent           # body proxy: resolves .part for new paragraphs
NORMAL = P[34].style          # a body-text style object
H2 = P[33].style              # Heading 2


def check(i, frag):
    assert frag in P[i].text, f"index {i} expected '{frag[:30]}' got '{P[i].text[:40]}'"


def set_text(p, text):
    if p.runs:
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
        p.runs[0].text = text
    else:
        p.add_run(text)


def repl(i, old, new):
    check(i, old)
    set_text(P[i], P[i].text.replace(old, new))


def add_cite(i, sentence):
    P[i].add_run(sentence)


def cit(*keys):
    """join LBU in-text cites into one parenthetical: (A, yr; B, yr)"""
    inner = "; ".join(REFS[k]["intext"][1:-1] for k in keys)
    return f" ({inner})"


# ---------------------------------------------------------------- 1. capture
# capture element refs up-front (identity stable through moves/deletes)
ch8_block = [P[i]._p for i in range(238, 259)]      # 238..258 (intro+RQ secs+figs+synth)
ch8_h1 = P[237]._p                                   # 'Chapter 8' heading -> delete
disc = P[192]._p                                     # '5.5 Discussion' -> move block before it
check(237, "Chapter 8"); check(192, "5.5 Discussion"); check(238, "foundational study")

# diagram anchors (first body paragraph of each target section)
DIAGS = [
    (110, "statechart.png", 5.4, "Figure 3.2: Fault-and-detection lifecycle for the early-detection "
        "study (RQ-D): a fault ramps the system from Normal through Degrading to an SLO Breach, while "
        "the proactive forecaster can raise an Early-Warning before the breach and the reactive "
        "classifier fires only once a window has breached."),
    (141, "usecase.png", 5.9, "Figure 4.1: Use-case view of the experiment harness. The researcher "
        "(and CI scheduler) drive the three primary studies — RQ-D early detection, RQ-A robustness and "
        "RQ-O cost-optimisation — each of which includes telemetry generation, observability "
        "configuration and evaluation."),
    (143, "architecture.png", 6.3, "Figure 4.2: System architecture organised around the three primary "
        "research questions. A single shared telemetry stream feeds configuration-aware feature "
        "extraction, which fans out into the RQ-D (early-detection), RQ-A (robustness) and RQ-O "
        "(cost-optimisation) experiment lanes before evaluation and reporting."),
    (153, "class.png", 6.3, "Figure 4.3: Class model of the detection pipeline and the three experiment "
        "types. Each experiment specialises an abstract Experiment that composes an ObsConfig, a "
        "BaselineModel and a Metrics object over Windows produced by the telemetry generator."),
    (156, "collaboration.png", 5.6, "Figure 4.4: Collaboration (communication) view of an RQ-O "
        "cost-optimisation run, in which the OptimiseExperiment coordinates the generator, feature "
        "builder, detector, metrics and Pareto analyser across the twenty observability configurations."),
    (160, "sequence.png", 6.3, "Figure 4.5: Sequence of an RQ-D early-detection run, contrasting the "
        "reactive classifier — which can only alarm once a window has breached the SLO — with the "
        "proactive forecaster, which projects the latency trend to warn before the breach."),
]
diag_anchors = [(P[i]._p, img, w, cap) for i, img, w, cap in DIAGS]

# ---------------------------------------------------------------- 2. foundational relabels
repl(165, "5.2 RQ1: The Effect of Observability Completeness on Detection",
     "5.2 Foundational Findings: Observability Completeness and Detection Accuracy")
repl(174, "5.3 RQ2: The Comparison of Algorithm Families",
     "5.3 Foundational Findings: Algorithm Families and Model Selection")
repl(183, "5.4 RQ3: The Contribution of Distributed Traces to Localisation",
     "5.4 Foundational Findings: Distributed Traces and Root-Cause Localisation")
repl(166, "The first research question asks how the completeness",
     "The foundational study first asks how the completeness")
repl(175, "The second research question asks which machine learning algorithms",
     "The foundational study next asks which machine learning algorithms")
repl(184, "The third research question asks what the specific contribution",
     "The foundational study finally asks what the specific contribution")

# ---------------------------------------------------------------- 3. reword residual old-RQ language
set_text(P[90],
    "These gaps are addressed in two stages. A foundational study first asks how the completeness of "
    "observability data affects the accuracy and timeliness of detection, addressing Gap 1 by "
    "constructing a controlled experiment in which the same machine learning pipeline is evaluated "
    "across four observability configurations of increasing richness; which machine learning algorithms "
    "and architectures are most effective on multimodal data, refining Gap 1 by recognising that the "
    "benefit of additional modalities may depend on the model's capacity to exploit them, and informed "
    "throughout by the requirements articulated by Raeiszadeh, Estrada-Solano and Glitho (2026); and the "
    "specific contribution of distributed traces to root cause localisation, responding to the "
    "multimodal root cause analysis evidence of Han et al. (2024) by isolating the trace modality and "
    "quantifying its marginal benefit. Building on these foundational findings, the dissertation then "
    "pursues its three primary research questions — RQ-D (timeliness), RQ-A (robustness) and RQ-O "
    "(cost-efficiency) — which extend the analysis from whether richer telemetry helps to how detection "
    "performs along the operational dimensions of lead time, resilience to telemetry degradation and "
    "cost" + cit("cheng2023aiops", "darban2025dltsad") + ".")

repl(95, "designed to answer the three research questions established in the preceding chapters",
     "designed to answer both the foundational study and the three primary research questions "
     "(RQ-D, RQ-A and RQ-O) established in the preceding chapters")
repl(125, "which is the subject of the first two research questions",
     "which underpins both the foundational study and the primary questions")
repl(125, "The second research question requires a comparison of algorithm families",
     "The foundational study requires a comparison of algorithm families")
repl(125, "For the localisation task, which is the subject of the third research question,",
     "For the localisation task, which forms part of the foundational study,")
repl(128, "the analysis of the third research question principally relies",
     "the foundational localisation analysis principally relies")
repl(151, "because it is the mechanism underlying the third research question",
     "because it is the mechanism underlying the foundational localisation analysis")
repl(153, "implements the several algorithm families required by the second research question",
     "implements the several algorithm families required by the foundational algorithm comparison")
repl(154, "The localisation component addresses the third research question.",
     "The localisation component addresses the foundational localisation analysis.")
repl(154, "which is precisely the quantity the third research question seeks",
     "which is precisely the quantity the foundational localisation analysis seeks")
repl(199, "The second research question's findings engage with",
     "The foundational algorithm-comparison findings engage with")

set_text(P[164],
    "This chapter presents and discusses the results of the experiment in two stages. It first reports "
    "the foundational findings, organised so that each can be mapped to the question it addresses: "
    "Section 5.2 reports the effect of observability completeness on detection; Section 5.3 reports the "
    "comparison of algorithm families on the richest configuration; and Section 5.4 reports the "
    "contribution of distributed traces to root cause localisation. Building on these foundations, "
    "Sections 5.5 to 5.7 then report the dissertation's three primary research questions — RQ-D on the "
    "timeliness of detection (Section 5.5), RQ-A on its robustness to telemetry degradation "
    "(Section 5.6) and RQ-O on the cost-optimal observability configuration (Section 5.7). Section 5.8 "
    "synthesises the three primary dimensions, and Section 5.9 draws all of the findings together, "
    "relates them to the research literature reviewed in Chapter 2, and considers their implications "
    "and limitations. The foundational results in Sections 5.2 to 5.4 were produced by the "
    "deterministic harness described in Chapter 4, executed over three hundred episodes comprising both "
    "fault and normal-operation periods, yielding a dataset of some ten thousand eight hundred labelled "
    "telemetry windows of which one hundred and eighty-six were fault episodes.")

# ---------------------------------------------------------------- 4. promote Ch8 -> Ch5 (renumber)
repl(239, "8.1 RQ-D: How Early Can Detection Warn Before an SLO Breach?",
     "5.5 RQ-D (Primary): How Early Can Detection Warn Before an SLO Breach?")
repl(245, "8.2 RQ-A: How Robust Is Detection to Telemetry Degradation?",
     "5.6 RQ-A (Primary): How Robust Is Detection to Telemetry Degradation?")
repl(251, "8.3 RQ-O: What Is the Cost-Optimal Observability Configuration?",
     "5.7 RQ-O (Primary): What Is the Cost-Optimal Observability Configuration?")
repl(257, "8.4 Synthesis", "5.8 Synthesis Across the Three Primary Dimensions")
repl(243, "Figure 8.1:", "Figure 5.4:")
repl(249, "Figure 8.2:", "Figure 5.5:")
repl(255, "Figure 8.3:", "Figure 5.6:")

# in-text citations in the promoted sections + bridge
add_cite(238, " The operational dimensions examined here — timeliness, robustness and cost — are exactly "
         "those emphasised in recent AIOps practice" + cit("cheng2023aiops", "somashekar2024gamma") + ".")
add_cite(244, " These results are consistent with recent work on proactive failure and SLA-breach "
         "prediction in cloud-native systems" +
         cit("lu2024misp", "fehresti2025sla", "guruge2025autoscaling", "darban2025dltsad") + ".")
add_cite(250, " The fragility of detection to degraded or missing telemetry echoes findings across the "
         "robustness literature" +
         cit("humblot2024noisy", "yuan2022trustworthy", "takahashi2024pu", "kumari2024investigation") + ".")
add_cite(256, " This cost-aware framing aligns with recent multi-objective and sampling-based approaches "
         "to efficient observability" +
         cit("grabher2025modeling", "sezgin2025multiobjective", "chen2024tracemesh", "fernando2025efficient") + ".")

# renumber the old Discussion -> 5.9
repl(192, "5.5 Discussion", "5.9 Discussion")
repl(193, "5.5.1 Analysis by Fault Category", "5.9.1 Analysis by Fault Category")
repl(197, "5.5.2 Relation to the Research Literature", "5.9.2 Relation to the Research Literature")
repl(201, "5.5.3 Practical Implications", "5.9.3 Practical Implications")
repl(203, "5.5.4 Threats to Validity", "5.9.4 Threats to Validity")

# move the Ch8 body block to just before the (now 5.9) Discussion, then drop the Ch8 heading
for el in ch8_block:
    disc.addprevious(el)
ch8_h1.getparent().remove(ch8_h1)

# ---------------------------------------------------------------- 5. embed diagrams
def insert_fig(anchor_p, img, width, caption):
    ip = OxmlElement("w:p"); anchor_p.addnext(ip)
    para = Paragraph(ip, BODY); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(DIAG / img), width=Inches(width))
    cp = OxmlElement("w:p"); ip.addnext(cp)
    cap = Paragraph(cp, BODY); cap.style = NORMAL; cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)

for anchor_p, img, w, cap in diag_anchors:
    insert_fig(anchor_p, img, w, cap)

# ---------------------------------------------------------------- 6. merge References (one A-Z list)
d2 = d  # alias
ref_head = None
for p in d2.paragraphs:
    if p.text.strip() == "References":
        ref_head = p
        break
# collect existing reference texts after the heading, then remove those paragraphs
existing = []
sib = ref_head._p.getnext()
to_remove = []
while sib is not None:
    if sib.tag.endswith("}p"):
        txt = "".join(t.text or "" for t in sib.iter() if t.tag.endswith("}t")).strip()
        if txt:
            existing.append(txt)
        to_remove.append(sib)
    sib = sib.getnext()
for el in to_remove:
    el.getparent().remove(el)

new_refs = [REFS[k]["full"] for k in REFS]
merged = existing + new_refs
def sortkey(s):
    s = s.lstrip()
    for art in ("A ", "An ", "The "):
        if s.startswith(art):
            s = s[len(art):]
            break
    return s.lower()
merged = sorted(set(merged), key=sortkey)

anchor = ref_head._p
for txt in merged:
    np = OxmlElement("w:p"); anchor.addnext(np); anchor = np
    par = Paragraph(np, BODY); par.style = NORMAL
    par.add_run(txt)

d2.save(OUT)

# ---------------------------------------------------------------- 7. verify
chk = docx.Document(OUT)
heads = [p.text.strip() for p in chk.paragraphs if (p.style and p.style.name == "Heading 2")
         and p.text.strip().startswith("5.")]
print("inline images:", len(chk.inline_shapes))
print("Ch5 sections:")
for h in heads:
    print("   ", h[:70])
print("Chapter 8 still present:", any("Chapter 8" in p.text for p in chk.paragraphs))
nref = sum(1 for p in chk.paragraphs if " (20" in p.text and "Available at" in p.text or
           (p.text.strip().endswith(".") and ") '" in p.text))
print("merged reference entries:", len(merged))
print("saved ->", OUT.name)
