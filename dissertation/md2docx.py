"""Convert draft_chapters.md -> draft_chapters.docx with Leeds Beckett formatting.
12 pt body, 1.5 line spacing, 1in margins (1.25in left for binding)."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\ResearchWithDrSatish\traceflix-platform\dissertation\draft_chapters.md"
OUT = r"D:\ResearchWithDrSatish\traceflix-platform\dissertation\draft_chapters.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

# Base style: 12 pt, 1.5 spacing
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def add_inline(p, text):
    tokens = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
    for tk in tokens:
        if not tk:
            continue
        if tk.startswith('**') and tk.endswith('**'):
            r = p.add_run(tk[2:-2]); r.bold = True
        elif tk.startswith('`') and tk.endswith('`'):
            r = p.add_run(tk[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif tk.startswith('*') and tk.endswith('*'):
            r = p.add_run(tk[1:-1]); r.italic = True
        else:
            p.add_run(tk)


def callout(paras):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    shade(cell, 'F2F6FC')
    cell.paragraphs[0].text = ''
    first = True
    for ln in paras:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, ln)
        first = False
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '8')
        e.set(qn('w:color'), '9DB8D2')
        borders.append(e)
    tblPr.append(borders)
    doc.add_paragraph()


with open(SRC, encoding='utf-8') as f:
    lines = f.readlines()

i, n = 0, len(lines)
while i < n:
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    if stripped == '---':
        doc.add_page_break(); i += 1; continue
    if stripped == '':
        i += 1; continue

    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1)); txt = m.group(2)
        h = doc.add_heading('', level=min(level, 4))
        add_inline(h, txt)
        i += 1; continue

    if stripped.startswith('>'):
        block = []
        while i < n and lines[i].strip().startswith('>'):
            t = re.sub(r'^>\s?', '', lines[i].strip())
            block.append(t)
            i += 1
        paras, cur = [], []
        for b in block:
            if b == '':
                if cur: paras.append(' '.join(cur)); cur = []
            else:
                cur.append(b)
        if cur: paras.append(' '.join(cur))
        callout(paras)
        continue

    if stripped.startswith('|'):
        tbl_lines = []
        while i < n and lines[i].strip().startswith('|'):
            tbl_lines.append(lines[i].strip()); i += 1
        rows = []
        for tl in tbl_lines:
            if re.match(r'^\|[\s:\-\|]+\|?$', tl):
                continue
            rows.append([c.strip() for c in tl.strip('|').split('|')])
        if rows:
            ncol = len(rows[0])
            table = doc.add_table(rows=0, cols=ncol)
            table.style = 'Light Grid Accent 1'
            for ri, r in enumerate(rows):
                cells = table.add_row().cells
                for ci in range(ncol):
                    val = r[ci] if ci < len(r) else ''
                    cells[ci].paragraphs[0].text = ''
                    pp = cells[ci].paragraphs[0]
                    add_inline(pp, val)
                    if ri == 0:
                        for rr in pp.runs: rr.bold = True
        doc.add_paragraph()
        continue

    if re.match(r'^[-*]\s+', stripped):
        while i < n and re.match(r'^[-*]\s+', lines[i].strip()):
            txt = re.sub(r'^[-*]\s+', '', lines[i].strip())
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, txt)
            i += 1
        continue

    if re.match(r'^\d+\.\s+', stripped):
        while i < n and re.match(r'^\d+\.\s+', lines[i].strip()):
            txt = re.sub(r'^\d+\.\s+', '', lines[i].strip())
            p = doc.add_paragraph(style='List Number')
            add_inline(p, txt)
            i += 1
        continue

    para = [line]; i += 1
    while i < n and lines[i].strip() != '' and not re.match(
            r'^(#{1,4}\s|>|\||[-*]\s|\d+\.\s|---)', lines[i].strip()):
        para.append(lines[i].rstrip('\n')); i += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, ' '.join(s.strip() for s in para))

doc.save(OUT)
print("saved", OUT)
